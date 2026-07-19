from pathlib import Path
from docx import Document
from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont

DOSSIER_SOURCE = Path("documents_a_trier")
DOSSIER_SOURCE.mkdir(exist_ok=True)


def creer_docx_test():
    """Cree un vrai fichier .docx avec du contenu client"""
    doc = Document()
    doc.add_paragraph("Convention d'honoraires")
    doc.add_paragraph("Pour le compte de notre client, Monsieur Julien Fabre,")
    doc.add_paragraph("concernant un litige commercial avec son fournisseur.")

    chemin = DOSSIER_SOURCE / "test_word.docx"
    doc.save(str(chemin))
    print(f"Cree : {chemin}")


def creer_pdf_test():
    """Cree un vrai fichier .pdf avec du texte reel (non scanne)"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, "Assignation en justice\n\nAffaire concernant notre cliente, Madame Claire Moreau,\ndans le cadre d'un litige de copropriete.")

    chemin = DOSSIER_SOURCE / "test_pdf.pdf"
    pdf.output(str(chemin))
    print(f"Cree : {chemin}")


def creer_pdf_scanne_test():
    """Cree un PDF qui simule un document scanne : une image sans aucun texte reel dedans"""
    image = Image.new("RGB", (800, 400), color="white")
    dessin = ImageDraw.Draw(image)

    texte = (
        "Contrat de bail commercial\n\n"
        "Client : Monsieur Thomas Girard\n"
        "Adresse du bien : 5 rue des Lilas"
    )

    try:
        police = ImageFont.truetype("arial.ttf", 24)
    except Exception:
        police = ImageFont.load_default()

    dessin.multiline_text((30, 30), texte, fill="black", font=police, spacing=10)

    chemin = DOSSIER_SOURCE / "test_pdf_scanne.pdf"
    image.save(str(chemin), "PDF")
    print(f"Cree (PDF scanne simule) : {chemin}")


if __name__ == "__main__":
    creer_docx_test()
    creer_pdf_test()
    creer_pdf_scanne_test()